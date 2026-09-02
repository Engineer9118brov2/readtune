from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).resolve().parents[1] / "docs" / "ReadTune-Award-Readiness-Plan.docx"

INK = "12212E"
NAVY = "164E63"
TEAL = "0F766E"
GOLD = "A16207"
RED = "B42318"
MIST = "E6F3F1"
PALE_BLUE = "EAF2F8"
PALE_GOLD = "FFF7E1"
LIGHT = "F4F6F8"
MUTED = "52616B"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for col, width in zip(grid.gridCol_lst, widths):
        col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_w = cell._tc.tcPr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.tcPr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def no_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "nil")
        borders.append(tag)
    tbl_pr.append(borders)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("ReadTune award readiness  |  ")
    run.font.name = "Aptos"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def add_text(paragraph, text, bold=False, color=INK, size=10.5, italic=False):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Aptos"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    return run


def add_body(doc, text, after=7):
    p = doc.add_paragraph(style="Body")
    p.paragraph_format.space_after = Pt(after)
    add_text(p, text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.16)
        add_text(p, item)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    add_text(p, text, bold=True, color=NAVY if level == 1 else TEAL, size=16 if level == 1 else 12.5)
    return p


def add_callout(doc, label, text, fill=MIST, label_color=TEAL):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    set_table_widths(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=150, start=200, bottom=150, end=200)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    add_text(p, label.upper() + "  ", bold=True, color=label_color, size=9)
    add_text(p, text, size=10)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_widths(table, widths)
    header = table.rows[0]
    for cell, label in zip(header.cells, headers):
        set_cell_shading(cell, NAVY)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        add_text(p, label, bold=True, color="FFFFFF", size=9)
    prevent_row_split(header)
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            set_cell_shading(cell, "FFFFFF" if row_index % 2 == 0 else LIGHT)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            add_text(p, value, size=9.2)
        prevent_row_split(table.rows[-1])
    doc.add_paragraph().paragraph_format.space_after = Pt(3)
    return table


def add_checklist(doc, title, items):
    add_heading(doc, title, 2)
    for item in items:
        p = doc.add_paragraph(style="Body")
        p.paragraph_format.space_after = Pt(4)
        add_text(p, "[ ] ", bold=True, color=TEAL)
        add_text(p, item)


def make_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12
    body = styles.add_style("Body", WD_STYLE_TYPE.PARAGRAPH)
    body.base_style = normal
    for level, size in ((1, 16), (2, 12.5), (3, 11)):
        style = styles[f"Heading {level}"]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(NAVY if level == 1 else TEAL)
        style.paragraph_format.space_before = Pt(14 if level == 1 else 10)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_text(header, "READTUNE  /  GATEWAYHACKS 2026", bold=True, color=TEAL, size=8.5)
    add_page_number(section.footer.paragraphs[0])

    # Cover
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(54)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(p, "AWARD READINESS PLAN", bold=True, color=TEAL, size=11)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(7)
    add_text(p, "ReadTune", bold=True, color=INK, size=32)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    add_text(p, "A 32-day plan to become a credible Accessibility & Health award contender", color=MUTED, size=15)
    add_callout(doc, "The decision", "Do not chase a nominal 100/100 with unsupported claims. Build a judging package that earns near-perfect credibility: a reliable live experience, real participant evidence, disciplined claims, and a technical story that survives questions.", PALE_BLUE, NAVY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(p, "Working window: August 31 - October 2, 2026", color=MUTED, size=10)
    add_text(p, "  |  Chrome Web Store review: reserve 3 days", color=MUTED, size=10)
    doc.add_page_break()

    add_heading(doc, "What a winning submission must prove")
    add_body(doc, "The build is already unusually ambitious for its age. The remaining month should turn capability into evidence. Judges should be able to see, within five minutes, that ReadTune solves a real reading-friction problem for real people without asking them to surrender their data.")
    add_table(doc, ["Rubric", "What must be true by October 2", "Proof the judge sees"], [
        ("Social impact (40)", "The product was shaped by people who experience reading friction; the team can show what changed because of them.", "One consented participant moment, 5-8 structured sessions, a change log, and early launch feedback."),
        ("Technical execution (30)", "Every advertised path works under demo conditions, including first-use failure handling and restoration.", "A rehearsed article, PDF, and TTS flow; automated checks; device smoke tests; an offline or fallback route."),
        ("Innovation (20)", "Personalization is framed honestly, and the on-device voice is real rather than a slideware claim.", "Calibration as a starting recommendation, plus opt-in Piper running locally after a one-time model download."),
        ("Design (10)", "The interface feels calm and intentional, and the product tells the truth at every decision point.", "One coherent visual system, clean settings language, and no high-confidence claims the evidence cannot bear."),
    ], [1500, 4300, 3560])
    add_callout(doc, "Priority order", "1) participant evidence and feedback-driven changes; 2) demo reliability; 3) Piper readiness and licensing; 4) calibration language and method; 5) launch feedback. A new feature belongs below these unless it directly fixes participant friction.", PALE_GOLD, GOLD)

    add_heading(doc, "The score pathway")
    add_body(doc, "A literal 100/100 is not controllable because judging is comparative. This plan targets the evidence that makes a 90-plus performance plausible against strong competition, especially in the two dimensions that dominate the rubric: social impact and technical execution.")
    add_table(doc, ["Current risk", "Intervention", "Expected judging effect"], [
        ("Impact feels inferred, not observed.", "Run 5-8 short, consented sessions with adults who self-identify as dyslexic, ADHD, or experience reading fatigue. Record problems, requests, and what changed.", "Moves impact from mission statement to demonstrable co-design."),
        ("Calibration sounds more certain than it is.", "Rename results as recommendations; show them as an optional starting point; make the limits explicit in the demo and copy.", "Prevents a statistics-minded judge from discounting the whole product."),
        ("Piper is easy to miss or risky to overpromise.", "Feature it only after first-use, offline-after-download, lower-end-device, and license checks pass. Otherwise demo browser TTS and describe Piper as opt-in beta work.", "Creates a memorable technical distinction without a brittle live gamble."),
        ("Web Store traction arrives late.", "Submit early, then use community outreach for feedback requests, not vanity installs. Keep a simple public changelog of feedback-to-change decisions.", "Makes launch feedback legible as learning, not marketing."),
    ], [2050, 4200, 3110])

    add_heading(doc, "Four-week operating plan")
    add_table(doc, ["Dates", "Outcome", "Non-negotiable work", "Exit test"], [
        ("Aug 31 - Sep 4", "Release readiness and founder demo", "Lock a known-good article and PDF; submit to the Chrome Web Store; prepare launch posts and a feedback intake route; record a baseline founder demo.", "Store package submitted; launch materials ready; every demo path works twice in a row."),
        ("Sep 5 - Sep 11", "Launch and post-launch recruiting", "Publish immediately after approval; post in relevant communities with an explicit feedback ask; recruit only from people who discover the published extension; run 2-3 sessions.", "At least two real launch users are scheduled or tested; a feedback log has initial observations."),
        ("Sep 12 - Sep 20", "Evidence-driven product changes", "Run 3-5 additional sessions with post-launch users; ship only recurring friction fixes; soften calibration and privacy copy; complete Piper readiness decision.", "At least two concrete changes trace directly to post-launch feedback; no unqualified outcome claims remain."),
        ("Sep 21 - Sep 27", "Proof and rehearsal", "Film the participant only after they have used the product; cut captioned proof clip; rehearse five-minute demo; test on a second machine or Chromebook if claiming broad support.", "Participant approves their clip; the demo survives a no-network or fallback rehearsal."),
        ("Sep 28 - Oct 2", "Submission lock", "Freeze feature work; run checks; finalize Devpost, video, store listing, and judge FAQ; rehearse hostile questions.", "No last-minute features; every headline claim has a source or a clear limitation."),
    ], [1250, 1850, 4000, 2260])
    add_callout(doc, "First 48 hours", "Submit the store package; lock the demo article and PDF; record a baseline founder demo in one take; and prepare one clear public feedback request. Recruitment starts only after the extension is publicly available.", PALE_GOLD, GOLD)
    doc.add_page_break()

    add_heading(doc, "Real-user evidence plan")
    add_body(doc, "Real users are the right evidence for impact. They do not turn the calibration into a clinical result; they show whether ReadTune is useful, respectful, and usable in an actual reading routine. Keep research lightweight, consented, and specific.")
    add_table(doc, ["Element", "Minimum standard", "What to capture"], [
        ("Participants", "5-8 adults recruited after launch from people who discover the published extension. Seek at least two relevant experiences, such as dyslexia, ADHD, or persistent reading fatigue.", "Self-described context, discovery channel, device type, reading task, prior assistive tools. Use participant IDs, not names, in the log."),
        ("Session", "20-30 minutes. Ask the participant to bring or choose a realistic article, assignment, handout, or PDF.", "Where they hesitate, what they choose, what they reject, whether they would use it again, and one direct quote if consented."),
        ("Compensation and consent", "Pay for time if possible. Let people stop, withdraw a clip, or refuse recording with no penalty.", "A signed or recorded consent decision for research notes and separately for public video/quotes."),
        ("Change log", "Make at least two product changes from repeated signals. Do not invent a change if feedback is mixed.", "Signal -> decision -> shipped change -> participant response after retest, if available."),
    ], [1550, 4400, 3160])
    add_callout(doc, "Recruiting script", "I built a free Chrome extension for people who experience reading friction. I am looking for adults who are willing to try it on one real reading task for 20-30 minutes. This is product feedback, not medical research. I will pay [amount] for your time; recording is optional, and you can stop at any time.", MIST, TEAL)
    doc.add_page_break()

    add_heading(doc, "Calibration: make it useful, not pseudo-scientific")
    add_body(doc, "The current calibration is an intelligent product heuristic, not a statistically defensible test. The winning move is not to hide that fact. Make the workflow humble and helpful: it offers a small set of settings to try, then lets the reader keep, change, or ignore them.")
    add_table(doc, ["Before October 2", "Change", "Reason"], [
        ("Copy", "Replace “what actually helps you read,” “confidence,” and “best setup” with “settings to try,” “how consistent this result was,” and “your starting setup.”", "The current study uses one short passage per condition, fixed baseline order, self-timed reading, and heuristic score weights."),
        ("Product", "Add an explicit “keep exploring” / “not for me” action on results. Let users see which settings were tested and which were not.", "Prevents the profile from claiming authority over TTS, size, tint, or other settings the test did not measure."),
        ("Demo", "Say: “This is a quick, local preference check, not a diagnostic or clinical assessment.”", "A judge who knows statistics will trust a product that states its boundary."),
        ("If time permits", "Counterbalance the baseline and repeat one condition, then describe it as a more robust pilot, not proof.", "Improves order-effect exposure but still does not create a validated assessment."),
    ], [2050, 4000, 3060])

    add_heading(doc, "Piper: a high-upside, gated differentiator")
    add_body(doc, "Piper can be one of the strongest technical and innovation moments: a natural-sounding voice that runs locally after an opt-in one-time model download. But it must never outrun its proof. The repository contains the implemented client, worker, controls, and test coverage; the planning document is behind the code and should be updated before submission.")
    add_checklist(doc, "Piper ship gate", [
        "Confirm end-to-end first use on a clean Chrome profile: permission, visible model-download progress, synthesis, sentence highlighting, and recovery from a failed download.",
        "Test on one constrained device or a meaningful CPU/network-throttle setup. If it cannot keep up, do not promise Chromebook performance.",
        "Resolve the licensing note in docs/PIPER.md before presenting or distributing the packaged runtime. The document itself flags an espeak-ng GPL concern; do not treat it as closed until verified.",
        "Make copy precise: the optional model downloads from Hugging Face once; reading text remains local. Do not use “nothing leaves your device” as a blanket statement.",
        "Keep browser TTS as the known-good fallback in the demo. Piper is a differentiator, not a single point of failure.",
    ])

    doc.add_page_break()
    add_heading(doc, "Five-minute judge demo")
    add_table(doc, ["Time", "Moment", "What it proves"], [
        ("0:00-0:25", "A real participant’s one-sentence problem statement or approved clip, followed by one realistic hard-to-read source page.", "Impact is grounded in a person, not an abstract persona."),
        ("0:25-1:20", "Reader View transforms the source. Let the participant’s chosen settings remain visible; do not claim they are universally better.", "The core transformation is immediate and legible."),
        ("1:20-2:05", "Read-aloud with sentence-following. Use Piper only if the ship gate is clear; otherwise use browser voice and mention local-first design.", "Accessible multimodal support works, not just a styling layer."),
        ("2:05-2:45", "Show the quick calibration as an optional starting recommendation, including the “you can change this” decision point.", "Personalization without medical or statistical overclaim."),
        ("2:45-3:35", "Open the same result on a PDF or in-page restyle. Demonstrate restoration or exit cleanly.", "Execution across the surfaces a reader actually encounters."),
        ("3:35-4:20", "Show the feedback-to-change log: two specific participant requests and the shipped responses.", "Social impact is a development practice, not a slogan."),
        ("4:20-5:00", "Close with the privacy boundary: local by default; optional services are clearly disclosed; no account or tracking needed.", "Trust and access are built into the product model."),
    ], [800, 5300, 3010])

    add_heading(doc, "Claim discipline")
    add_table(doc, ["Use this", "Avoid this", "Why"], [
        ("“A free, local-first reading-support extension that lets readers try a calmer setup across articles and PDFs.”", "“ReadTune fixes dyslexia” or “treats reading disorders.”", "The product supports reading; it does not provide a medical intervention."),
        ("“A quick preference check that suggests settings to try.”", "“Measures what actually helps you read” or “your best setup.”", "The calibration is a heuristic with substantial design shortcuts."),
        ("“Reading text stays on-device with browser TTS and Piper; optional ElevenLabs sends selected text to the user’s own account.”", "“Nothing leaves your device.”", "The blanket claim is false for opted-in online voice and the Piper model download."),
        ("“Participant feedback shaped these changes.”", "“Proven to improve comprehension.”", "Usability sessions are excellent impact evidence but do not establish clinical efficacy."),
    ], [2880, 2880, 3350])

    doc.add_page_break()
    add_heading(doc, "Submission lock checklist")
    add_checklist(doc, "Product", [
        "Known-good demo article and PDF are locally saved or reliably available.",
        "Reader View, in-page restyle, PDF view, calibration, and browser TTS pass twice in a fresh profile.",
        "If Piper is shown: every ship-gate item passes. If not: remove it from hero claims and show it only as clearly labelled optional work.",
        "All privacy and store copy distinguish browser TTS, Piper’s one-time download, and optional ElevenLabs transmission.",
    ])
    add_checklist(doc, "Evidence", [
        "5-8 sessions are documented with consent, participant IDs, observations, and decisions.",
        "At least two feedback-driven changes are shipped and visible in a small changelog.",
        "The demo participant has reviewed and approved the exact public clip or quote.",
        "Any launch feedback is dated and described as early feedback, not product efficacy evidence.",
    ])
    add_checklist(doc, "Presentation", [
        "The five-minute demo is rehearsed enough to run without narration improvisation.",
        "The team can answer: What does calibration measure? What leaves the device? What did users ask you to change? What happens when a site is complex?", 
        "Every headline in Devpost, the store listing, popup, results screen, and marketing site survives the claim-discipline table above.",
    ])
    add_callout(doc, "Final standard", "The credible version of this story is stronger than a perfect-sounding one: ReadTune was built quickly, then tested respectfully with readers, improved from what they said, and delivered as a free, local-first tool with honest limits. That is the story a strong Accessibility & Health judge can defend in the room.", PALE_BLUE, NAVY)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    make_doc()
